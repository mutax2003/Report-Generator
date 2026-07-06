"""
Generate Phase I appendix Word templates (D drilling waste checklist, G calc tables).

Run: python scripts/create_appendix_templates.py
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from engine import ECOVENTURE_CONSULTANT, _add_docx_table_loop  # noqa: E402

APPENDIX_DIR = ROOT / "samples" / "appendices"
APPENDIX_A = APPENDIX_DIR / "appendix_a_qp_declaration.docx"
APPENDIX_D = APPENDIX_DIR / "appendix_d_drilling_waste_checklist.docx"
APPENDIX_G = APPENDIX_DIR / "appendix_g_waste_calc_tables.docx"


def generate_appendix_a_template(path: Path) -> None:
    """Appendix A — QP professional declaration (R&R / SED 002)."""
    doc = Document()
    doc.add_heading("Appendix A", level=0)
    doc.add_heading("Professional Declaration", level=1)
    doc.add_paragraph("Prepared for: {{ client_name }}")
    doc.add_paragraph("Well: {{ well_name }}")
    doc.add_paragraph("UWI: {{ uwi }}")
    doc.add_paragraph("Project: {{ project_number }}")
    doc.add_paragraph("")
    doc.add_paragraph("Consultant: {{ consultant_name }}")
    doc.add_paragraph("Company: {{ company }}")
    doc.add_paragraph("Qualified person(s): {{ qp_names }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }}")
    doc.add_paragraph("Date of issue: {{ date_of_issue }}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "I confirm that I have prepared or supervised the preparation of this "
        "Phase 1 Environmental Site Assessment report and that the information "
        "contained herein is accurate to the best of my knowledge."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Signature: _________________________  Date: _____________")
    doc.add_paragraph("Name (print): {{ prepared_by }}")
    doc.add_paragraph("Professional designation: _________________________")
    doc.add_paragraph(f"{{{{ consultant_name }}}} ({ECOVENTURE_CONSULTANT})")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def generate_appendix_d_template(path: Path) -> None:
    """Appendix D — drilling waste notification and Option 1/2 compliance checklist."""
    doc = Document()
    doc.add_heading("Appendix D", level=0)
    doc.add_heading("Drilling Waste Notification and Compliance Checklist", level=1)
    doc.add_paragraph("Prepared for: {{ client_name }}")
    doc.add_paragraph("Well: {{ well_name }}")
    doc.add_paragraph("UWI: {{ uwi }}")
    doc.add_paragraph("Project: {{ project_number }}")
    doc.add_paragraph("")
    doc.add_paragraph("Consultant: {{ consultant_name }}")
    doc.add_paragraph("Qualified person(s): {{ qp_names }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }}")
    doc.add_paragraph("Date: {{ date_of_issue }}")
    doc.add_paragraph("")
    doc.add_paragraph("AER drilling waste compliance option: {{ aer_waste_compliance_option }}")
    doc.add_paragraph("No drilling waste on site: {{ no_drilling_waste_on_site }}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Drilling waste disposal checklist items (AER Directive 050 / SED 002 section 10.4). "
        "Compliance option: {{ aer_waste_compliance_option }}."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Summary narrative:")
    doc.add_paragraph("{{ drilling_waste_summary }}")
    doc.add_paragraph("")
    doc.add_paragraph("Guideline scope: {{ dwda_guideline_summary }}")
    doc.add_paragraph("Compliance summary: {{ dwda_compliance_summary }}")
    doc.add_paragraph("")
    _add_docx_table_loop(
        doc,
        "DWDA compliance checklist (companion to official AER form):",
        "dwda_checklist_results",
        ["Section", "Item", "Status", "Response", "Detail"],
        ["section", "label", "status", "response", "detail"],
    )
    doc.add_paragraph("")
    _add_docx_table_loop(
        doc,
        "Drilling waste disposal events (checklist reference):",
        "drilling_waste",
        [
            "Mud type",
            "Volume (m3)",
            "Disposal type",
            "Method",
            "Location",
            "GPS",
            "Manifests",
        ],
        [
            "mud_type",
            "volume_m3",
            "disposal_type",
            "disposal_method",
            "location",
            "gps_coordinates",
            "waste_manifest_refs",
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "QP declaration: I confirm the drilling waste information above is accurate "
        "to the best of my knowledge for this Phase 1 ESA."
    )
    doc.add_paragraph("Signature: _________________________  Date: _____________")
    doc.add_paragraph(f"{{{{ consultant_name }}}} ({ECOVENTURE_CONSULTANT})")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def generate_appendix_g_template(path: Path) -> None:
    """Appendix G — drilling waste calculation tables."""
    doc = Document()
    doc.add_heading("Appendix G", level=0)
    doc.add_heading("Drilling Waste Calculation Tables", level=1)
    doc.add_paragraph("Prepared for: {{ client_name }}")
    doc.add_paragraph("Well: {{ well_name }} | UWI: {{ uwi }}")
    doc.add_paragraph("Compliance option: {{ aer_waste_compliance_option }}")
    doc.add_paragraph("Checklist scope: {{ dwda_checklist_scope }}")
    doc.add_paragraph("Report date: {{ date_of_issue }}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Salinity within each DWDA may be assessed using Equivalent Salinity Guidelines. "
        "Hydrocarbons, metals, and other parameters must meet Alberta Tier 1/2 within the "
        "DWDA and across the remainder of the lease."
    )
    doc.add_paragraph("")
    _add_docx_table_loop(
        doc,
        "Drilling waste calculation table:",
        "drilling_waste",
        [
            "Mud type",
            "Volume (m3)",
            "Disposal type",
            "Method",
            "Location",
            "GPS",
            "Sump depth (m)",
            "Cover (m)",
            "Remote cert #",
            "Manifests",
            "DWDA ID",
            "Area (m2)",
            "Salinity exceedance",
        ],
        [
            "mud_type",
            "volume_m3",
            "disposal_type",
            "disposal_method",
            "location",
            "gps_coordinates",
            "sump_depth_m",
            "cover_depth_m",
            "remote_cert_number",
            "waste_manifest_refs",
            "dwda_id",
            "area_m2",
            "salinity_exceedance",
        ],
    )
    doc.add_paragraph("")
    doc.add_heading("DWDA calculation summary", level=2)
    doc.add_paragraph("{{ dwda_calc_summary }}")
    doc.add_paragraph(
        "Metal (barite sacks per metre): {{ dwda_metal_sacks_per_metre }} — "
        "Pass: {{ dwda_metal_pass }}"
    )
    doc.add_paragraph(
        "Salt (NaOH-equiv sacks per m³): {{ dwda_salt_sacks_per_m3 }} — "
        "Pass: {{ dwda_salt_pass }}"
    )
    doc.add_paragraph(
        "DST resistivity sacks: {{ dwda_dst_resistivity_sacks_total }} | "
        "Chloride sacks: {{ dwda_dst_chloride_sacks_total }}"
    )
    _add_docx_table_loop(
        doc,
        "Calculation detail:",
        "dwda_calculations",
        ["Type", "Result", "Objective", "Pass", "Notes"],
        ["calc_type", "result_value", "objective", "pass", "notes"],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Export this appendix to PDF before OneStop submission when required."
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    generate_appendix_a_template(APPENDIX_A)
    generate_appendix_d_template(APPENDIX_D)
    generate_appendix_g_template(APPENDIX_G)
    print(f"Wrote: {APPENDIX_A}")
    print(f"Wrote: {APPENDIX_D}")
    print(f"Wrote: {APPENDIX_G}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
