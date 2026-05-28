"""AI feature tests (offline / heuristic paths)."""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from ai.consistency import check_consistency
from ai.copilot import explain_preflight
from ai.exceedance_notes import notes_for_lab_rows
from ai.lab_extract import _heuristic_rows
from ai.narrative import draft_narratives
from ai.rag import load_corpus, retrieve
from ai.template_tagger import suggest_template_tags, suggestions_to_markdown
from engine import generate_sample_template_docx
from template_tools import run_preflight

ROOT = Path(__file__).resolve().parents[1]
SAMPLES = ROOT / "samples"


class AiFeatureTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        xlsx = SAMPLES / "sample_data.xlsx"
        docx = SAMPLES / "sample_template.docx"
        if not xlsx.is_file() or not docx.is_file():
            from engine import generate_sample_excel, generate_sample_template_docx

            SAMPLES.mkdir(parents=True, exist_ok=True)
            generate_sample_excel(str(xlsx))
            generate_sample_template_docx(str(docx))
        cls.excel = (SAMPLES / "sample_data.xlsx").read_bytes()
        cls.template = (SAMPLES / "sample_template.docx").read_bytes()

    def test_template_tagger_rules(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            path = tmp.name
            from docx import Document

            doc = Document()
            doc.add_paragraph("Client: [Company]")
            doc.add_paragraph("Site: Client Full Name")
            doc.save(path)
            data = Path(path).read_bytes()
        try:
            suggestions, audit = suggest_template_tags(data, use_llm=False)
            tags = {s.jinja_tag for s in suggestions}
            self.assertTrue(any("company" in t for t in tags))
            md = suggestions_to_markdown(suggestions)
            self.assertIn("Template tagging", md)
            self.assertFalse(audit.used_llm)
        finally:
            Path(path).unlink(missing_ok=True)

    def test_heuristic_lab_rows(self) -> None:
        text = "Benzene 0.8 ug/L 5.0\nTCE 12.5 ug/L 5.0"
        rows = _heuristic_rows(text)
        self.assertGreaterEqual(len(rows), 1)

    def test_rag_corpus(self) -> None:
        chunks = load_corpus()
        self.assertGreaterEqual(len(chunks), 1)
        hits = retrieve("Phase 2 environmental site assessment exceedance")
        self.assertGreaterEqual(len(hits), 1)

    def test_narrative_offline(self) -> None:
        ctx = {
            "site_name": "123 Road",
            "client_name": "Demo",
            "report_phase": "Phase 2",
            "lab_results": [{"exceedance_flag": "Y", "analyte": "TCE"}],
        }
        drafts, audit = draft_narratives(ctx, use_llm=False)
        self.assertGreaterEqual(len(drafts), 1)
        self.assertFalse(audit.used_llm)

    def test_narrative_phase1_ecoventure(self) -> None:
        from ai.narrative import sections_for_phase

        ctx = {
            "client_name": "Example Energy Ltd.",
            "consultant_name": "Ecoventure Inc.",
            "well_name": "Example 4D Windy 4-4-49-4",
            "report_phase": "Phase 1",
            "executive_summary": "Ecoventure Inc. prepared this Phase I ESA.",
            "drilling_waste_summary": "Option 1 disposal.",
        }
        self.assertIn("drilling_waste", sections_for_phase("Phase 1"))
        self.assertIn(
            "hydrogeologic_setting",
            sections_for_phase("Phase 1", "groundwater_monitoring"),
        )
        drafts, audit = draft_narratives(ctx, use_llm=False)
        self.assertGreaterEqual(len(drafts), 3)
        exec_d = next(d for d in drafts if d.section == "executive_summary")
        self.assertIn("Ecoventure", exec_d.text)
        self.assertFalse(audit.used_llm)

    def test_narrative_phase1_generates_signum_style_when_empty(self) -> None:
        from ai.narrative import draft_narratives

        ctx = {
            "client_name": "Example Energy Ltd.",
            "consultant_name": "Ecoventure Inc.",
            "well_name": "Example 4D Windy 4-4-49-4",
            "report_phase": "Phase 1",
            "spud_date": "15-Mar-2004",
            "drilling_waste_summary": "110 m3 disposed via LWD",
            "site_visit_completed": "No",
        }
        drafts, _ = draft_narratives(ctx, use_llm=False)
        exec_d = next(d for d in drafts if d.section == "executive_summary")
        self.assertIn("was contracted by", exec_d.text)
        self.assertIn("After reviewing regulatory information", exec_d.text)

    def test_copilot_offline(self) -> None:
        pf = run_preflight(
            self.excel,
            self.template,
            {"report_phase": "Phase 2", "date_of_issue": "2026-01-01"},
        )
        advice, audit = explain_preflight(
            pf, {"report_phase": "Phase 2"}, use_llm=False
        )
        self.assertTrue(advice.summary)
        self.assertFalse(audit.used_llm)

    def test_consistency_duplicate_analyte(self) -> None:
        ctx = {
            "site_name": "A",
            "site_address": "B far away",
            "lab_results": [
                {"analyte": "Benzene", "result": "1", "criteria": "5", "exceedance_flag": "N"},
                {"analyte": "Benzene", "result": "2", "criteria": "5", "exceedance_flag": "N"},
            ],
        }
        findings, _ = check_consistency(ctx, use_llm=False)
        codes = {f.code for f in findings}
        self.assertIn("duplicate_analyte", codes)

    def test_exceedance_notes_offline(self) -> None:
        lab = [
            {
                "analyte": "TCE",
                "result": "12",
                "unit": "ug/L",
                "criteria": "5",
                "exceedance_flag": "Y",
            }
        ]
        notes, audit = notes_for_lab_rows(lab, use_llm=False)
        self.assertEqual(len(notes), 1)
        self.assertIn("exceed", notes[0].note.lower())


if __name__ == "__main__":
    unittest.main()
