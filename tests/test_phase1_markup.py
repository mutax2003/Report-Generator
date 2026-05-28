"""Phase 1 PDF markup + AI integration tests."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from ai.template_tagger import suggest_template_tags
from phase1_markup import apply_tag_suggestions, enhance_phase1_markup
from phase1_pdf_text import Phase1PdfMeta, build_mvp_tag_replacements
from report_profile import get_recommended_fields


class TestPhase1Markup(unittest.TestCase):
    def test_phase1_allowed_keys_in_tagger(self) -> None:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            path = Path(tmp.name)
            doc = Document()
            doc.add_paragraph("Ecoventure Inc.")
            doc.add_paragraph("Phase 1 Environmental Site Assessment")
            doc.save(str(path))
            data = path.read_bytes()
        try:
            suggestions, audit = suggest_template_tags(
                data, use_llm=False, report_type="phase1_alberta"
            )
            tags = {s.jinja_tag for s in suggestions}
            self.assertTrue(any("company" in t for t in tags))
            self.assertIn("phase1_alberta", audit.features)
        finally:
            path.unlink(missing_ok=True)

    def test_mvp_replacements(self) -> None:
        meta = Phase1PdfMeta(
            project_number="260109R",
            client_name="Caltex Trilogy Inc.",
            well_name="100/07-34-055-02 W4M",
            site_name="100/07-34-055-02 W4M",
            uwi="100/07-34-055-02 W4M",
            report_title="2025 Phase 1 Environmental Site Assessment",
            report_month_year="January 9, 2026",
            lsd_from_filename="16-34-055-02W4M",
        )
        repl = build_mvp_tag_replacements(meta)
        self.assertIn("Caltex Trilogy Inc.", repl)
        self.assertEqual(repl["Caltex Trilogy Inc."], "{{ client_name }}")

    def test_enhance_minimal_docx(self) -> None:
        doc = Document()
        doc.add_paragraph("Caltex Trilogy Inc.")
        doc.add_paragraph("Ecoventure Inc.")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            path = Path(tmp.name)
            doc.save(str(path))
            raw = path.read_bytes()
        try:
            meta = Phase1PdfMeta(
                project_number="260109R",
                client_name="Caltex Trilogy Inc.",
                well_name="Site",
                site_name="Site",
                uwi="00/01-01-001-01W4/0",
                report_title="Phase I ESA",
                report_month_year="Jan 2026",
                lsd_from_filename="",
            )
            result = enhance_phase1_markup(
                raw, meta, use_ai=True, apply_ai_tags=True, use_llm=False
            )
            self.assertGreater(result.mvp_replacements, 0)
            from template_tools import scan_template

            scan = scan_template(result.docx_bytes)
            self.assertIn("client_name", scan.root_vars)
        finally:
            path.unlink(missing_ok=True)

    def test_profile_fields_nonempty(self) -> None:
        fields = get_recommended_fields("phase1_alberta")
        self.assertIn("client_name", fields)
        self.assertIn("executive_summary", fields)


if __name__ == "__main__":
    unittest.main()
