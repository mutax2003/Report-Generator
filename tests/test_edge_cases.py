"""Edge-case and regression tests for engine, security, and render paths."""

from __future__ import annotations

import io
import sys
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from engine import (  # noqa: E402
    LAB_SHEET,
    PROJECT_SHEET,
    ReportEngine,
    _lab_frame_to_records,
    _norm_key,
    _numeric_compare_exceeds,
    _project_row_to_dict,
    _truthy_exceedance,
    collect_template_root_vars,
    suggested_download_name,
    generate_sample_template_docx,
)
from security import (  # noqa: E402
    MAX_EXCEL_BYTES,
    MAX_LAB_ROWS,
    MAX_META_VALUE_LEN,
    SecurityError,
    clamp_context,
    sanitize_download_filename,
    sanitize_meta,
    validate_excel_upload,
    validate_template_upload,
)


def _minimal_xlsx_zip() -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Override PartName="/xl/workbook.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
            "</Types>",
        )
        zf.writestr("xl/workbook.xml", "<workbook/>")
    return buf.getvalue()


def _minimal_docx_bytes(extra_paragraph: str | None = None) -> bytes:
    doc = Document()
    doc.add_paragraph("Site: {{ site_name }}")
    if extra_paragraph:
        doc.add_paragraph(extra_paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _workbook_bytes(
    project: pd.DataFrame | None = None,
    lab: pd.DataFrame | None = None,
    *,
    include_project: bool = True,
    include_lab: bool = True,
    extra_sheets: dict[str, pd.DataFrame] | None = None,
) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as w:
        if include_project and project is not None:
            project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        if include_lab and lab is not None:
            lab.to_excel(w, sheet_name=LAB_SHEET, index=False)
        if extra_sheets:
            for name, df in extra_sheets.items():
                df.to_excel(w, sheet_name=name, index=False)
    return bio.getvalue()


class TestExcelEdgeCases(unittest.TestCase):
    def setUp(self) -> None:
        self.tpl = ROOT / "samples" / "sample_template.docx"
        if not self.tpl.is_file():
            self.tpl.parent.mkdir(parents=True, exist_ok=True)
            generate_sample_template_docx(str(self.tpl))
        self.template_bytes = self.tpl.read_bytes()

    def test_missing_project_sheet(self) -> None:
        xlsx = _workbook_bytes(
            project=pd.DataFrame({"x": [1]}),
            include_project=False,
            extra_sheets={"Wrong": pd.DataFrame({"a": [1]})},
        )
        engine = ReportEngine(xlsx, self.template_bytes)
        with self.assertRaises(ValueError) as ctx:
            engine.build_context({"report_phase": "Phase 2"})
        self.assertIn("ProjectData", str(ctx.exception))

    def test_missing_lab_sheet_phase2(self) -> None:
        xlsx = _workbook_bytes(
            project=pd.DataFrame({"site_name": ["A"]}),
            include_lab=False,
        )
        engine = ReportEngine(xlsx, self.template_bytes)
        with self.assertRaises(ValueError) as ctx:
            engine.build_context({"report_phase": "Phase 2"})
        self.assertIn("LabResults", str(ctx.exception))

    def test_phase1_without_lab_sheet(self) -> None:
        xlsx = _workbook_bytes(
            project=pd.DataFrame({"site_name": ["Phase1 Site"], "client_name": ["C"]}),
            include_lab=False,
        )
        engine = ReportEngine(xlsx, self.template_bytes)
        ctx = engine.build_context({"report_phase": "Phase 1"})
        self.assertEqual(ctx["site_name"], "Phase1 Site")
        self.assertEqual(ctx["lab_results"], [])

    def test_empty_project_data_rows(self) -> None:
        xlsx = _workbook_bytes(
            project=pd.DataFrame(columns=["site_name", "client_name"]),
            lab=pd.DataFrame(
                [{"Analyte": "X", "Result": 1, "Unit": "mg/L", "Exceedance": "N"}]
            ),
        )
        engine = ReportEngine(xlsx, self.template_bytes)
        with self.assertRaises(ValueError) as ctx:
            engine.build_context({"report_phase": "Phase 2"})
        self.assertIn("no data rows", str(ctx.exception))

    def test_empty_lab_results(self) -> None:
        xlsx = _workbook_bytes(
            project=pd.DataFrame({"site_name": ["Site"], "client_name": ["C"]}),
            lab=pd.DataFrame(columns=["Analyte", "Result", "Unit", "Exceedance"]),
        )
        engine = ReportEngine(xlsx, self.template_bytes)
        ctx = engine.build_context({"report_phase": "Phase 2"})
        self.assertEqual(ctx["lab_results"], [])

    def test_header_normalization(self) -> None:
        df = pd.DataFrame([{"Site Name": "ABC", "Client  Name": "Co"}])
        out = _project_row_to_dict(df)
        self.assertEqual(out.get("site_name"), "ABC")
        self.assertEqual(out.get("client_name"), "Co")

    def test_alternate_lab_column_names(self) -> None:
        df = pd.DataFrame(
            [
                {
                    "Parameter": "Lead",
                    "Value": 10,
                    "Units": "mg/kg",
                    "Standard": 5,
                    "Flag": "Y",
                }
            ]
        )
        rows = _lab_frame_to_records(df)
        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["analyte"], "Lead")
        self.assertEqual(rows[0]["exceedance_flag"], "Yes")

    def test_numeric_exceedance_without_flag(self) -> None:
        df = pd.DataFrame(
            [{"Analyte": "Cu", "Result": 10, "Unit": "mg/L", "Criteria": 5, "Exceedance": ""}]
        )
        rows = _lab_frame_to_records(df)
        self.assertEqual(rows[0]["exceedance_flag"], "Yes")

    def test_nan_cells_become_empty(self) -> None:
        df = pd.DataFrame([{"site_name": float("nan"), "client_name": "OK"}])
        out = _project_row_to_dict(df)
        self.assertEqual(out.get("site_name"), "")
        self.assertEqual(out.get("client_name"), "OK")


class TestExceedanceFlags(unittest.TestCase):
    def test_truthy_variants(self) -> None:
        for val in ("Y", "yes", "TRUE", "1", "x", "EXCEEDANCE", "exc"):
            self.assertTrue(_truthy_exceedance(val), msg=val)

    def test_non_exceedance(self) -> None:
        for val in ("N", "no", "", None, "0", "false"):
            self.assertFalse(_truthy_exceedance(val), msg=repr(val))

    def test_numeric_compare(self) -> None:
        self.assertTrue(_numeric_compare_exceeds(10, 5))
        self.assertFalse(_numeric_compare_exceeds(5, 10))
        self.assertFalse(_numeric_compare_exceeds("n/a", 5))


class TestRenderEdgeCases(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.tpl_path = ROOT / "samples" / "sample_template.docx"
        if not cls.tpl_path.is_file():
            raise unittest.SkipTest("Run scripts/create_samples.py first")
        cls.template_bytes = cls.tpl_path.read_bytes()

    def test_missing_template_var_warning(self) -> None:
        tpl = _minimal_docx_bytes("Extra: {{ not_in_excel }}")
        xlsx = _workbook_bytes(
            project=pd.DataFrame({"site_name": ["S"], "client_name": ["C"]}),
            lab=pd.DataFrame(
                [{"Analyte": "A", "Result": 1, "Unit": "u", "Exceedance": "N"}]
            ),
        )
        engine = ReportEngine(xlsx, tpl)
        _docx, warnings, _ctx, _rec = engine.render(
            meta={"report_phase": "Phase 2", "prepared_by": "T", "date_of_issue": "2026-01-01"}
        )
        self.assertTrue(any("not_in_excel" in w for w in warnings))

    def test_unicode_download_filename(self) -> None:
        name = suggested_download_name(
            {"site_name": "Site / Résumé 中文"},
            {"report_phase": "Phase 2", "date_of_issue": "2026-05-19"},
        )
        self.assertTrue(name.endswith(".docx"))
        self.assertNotIn("/", name)
        self.assertLessEqual(len(name), 200)

    def test_meta_overrides_excel_key(self) -> None:
        xlsx = _workbook_bytes(
            project=pd.DataFrame({"prepared_by": ["From Excel"], "site_name": ["S"]}),
            lab=pd.DataFrame(
                [{"Analyte": "A", "Result": 1, "Unit": "u", "Exceedance": "N"}]
            ),
        )
        engine = ReportEngine(xlsx, self.template_bytes)
        ctx = engine.build_context(
            {"prepared_by": "From Sidebar", "report_phase": "Phase 2"}
        )
        self.assertEqual(ctx["prepared_by"], "From Sidebar")

    def test_render_succeeds_with_minimal_fields(self) -> None:
        xlsx = _workbook_bytes(
            project=pd.DataFrame({"site_name": ["Only Site"]}),
            lab=pd.DataFrame(
                [{"Analyte": "A", "Result": 1, "Unit": "u", "Exceedance": "N"}]
            ),
        )
        engine = ReportEngine(xlsx, self.template_bytes)
        docx_bytes, warnings, ctx, _rec = engine.render(
            meta={"report_phase": "Phase 2", "date_of_issue": "2026-01-01"}
        )
        self.assertGreater(len(docx_bytes), 500)
        self.assertIsInstance(warnings, list)
        self.assertEqual(ctx["site_name"], "Only Site")


class TestSecurityEdgeCases(unittest.TestCase):
    def test_docx_rejected_as_excel(self) -> None:
        docx = _minimal_docx_bytes()
        with self.assertRaises(SecurityError):
            validate_excel_upload(docx)

    def test_xlsx_rejected_as_docx(self) -> None:
        xlsx = _minimal_xlsx_zip()
        with self.assertRaises(SecurityError):
            validate_template_upload(xlsx)

    def test_excel_size_limit(self) -> None:
        data = _minimal_xlsx_zip()
        with patch("security.MAX_EXCEL_BYTES", len(data) - 1):
            with self.assertRaises(SecurityError) as ctx:
                validate_excel_upload(data)
        self.assertIn("too large", str(ctx.exception))

    def test_sanitize_meta_truncates_long_values(self) -> None:
        long = "x" * (MAX_META_VALUE_LEN + 50)
        out = sanitize_meta({"prepared_by": long})
        self.assertEqual(len(out["prepared_by"]), MAX_META_VALUE_LEN)

    def test_sanitize_meta_skips_blank_keys(self) -> None:
        out = sanitize_meta({"": "x", "   ": "y", "ok": "z"})
        self.assertIn("ok", out)
        self.assertEqual(len(out), 1)

    def test_clamp_context_truncates_lab_rows(self) -> None:
        huge = [{"analyte": str(i)} for i in range(MAX_LAB_ROWS + 5)]
        ctx, warnings = clamp_context({"lab_results": huge})
        self.assertEqual(len(ctx["lab_results"]), MAX_LAB_ROWS)
        self.assertTrue(any("truncated" in w.lower() for w in warnings))

    def test_download_filename_unicode_and_empty(self) -> None:
        self.assertTrue(sanitize_download_filename("").endswith(".docx"))
        name = sanitize_download_filename("___\x00///.docx")
        self.assertTrue(name.endswith(".docx"))


class TestNormKey(unittest.TestCase):
    def test_spaces_and_case(self) -> None:
        self.assertEqual(_norm_key("  Site Name "), "site_name")

    def test_collect_vars_ignores_item_fields(self) -> None:
        tpl = ROOT / "samples" / "sample_template.docx"
        if not tpl.is_file():
            self.skipTest("no sample template")
        roots = collect_template_root_vars(tpl.read_bytes())
        self.assertNotIn("item", roots)


if __name__ == "__main__":
    unittest.main()
