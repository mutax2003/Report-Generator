"""
Document rendering for ESA reports: Excel -> Jinja2 context -> docxtpl.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any

import pandas as pd
from docx import Document
from docxtpl import DocxTemplate, RichText
from jinja2 import StrictUndefined
from jinja2.exceptions import TemplateError
from jinja2.sandbox import SandboxedEnvironment

from phase1_narrative import build_phase1_executive_summary
from report_profile import (
    ReportRuntimeConfig,
    list_keys_from_context,
    resolve_report_config,
    table_row_counts,
)

from security import (
    MAX_BATCH_REPORTS,
    MAX_LAB_ROWS,
    MAX_PROJECT_COLUMNS,
    MAX_PROJECT_ROWS,
    SecurityError,
    clamp_context,
    sanitize_download_filename,
    sanitize_meta,
    validate_excel_upload,
    validate_rendered_output,
    validate_template_upload,
    validation_bypass_enabled,
)


PROJECT_SHEET = "ProjectData"
LAB_SHEET = "LabResults"
GROUNDWATER_LAB_SHEET = "GroundwaterLab"
MONITORING_WELLS_SHEET = "MonitoringWells"
WATER_LEVELS_SHEET = "WaterLevels"
APECS_SHEET = "Apecs"
DRILLING_WASTE_SHEET = "DrillingWaste"
STORAGE_TANKS_SHEET = "StorageTanks"
DWDA_CHECKLIST_SHEET = "DwdaChecklist"
DWDA_CALCULATIONS_SHEET = "DwdaCalculations"

ECOVENTURE_CONSULTANT = "Ecoventure Inc."


def _s(value: Any) -> str:
    if value is None:
        return ""
    t = str(value).strip()
    if t.lower() in ("nan", "none"):
        return ""
    return t


def _norm_key(name: str) -> str:
    s = str(name).strip()
    s = re.sub(r"\s+", "_", s)
    return s.lower()


def _merge_dwda_calc_sheet(ctx: dict[str, Any]) -> None:
    """Promote DwdaCalculations sheet row into scalar context for calc engine."""
    from ecoventure_workbook import (
        cell_contract_provenance,
        flat_calc_row_from_sheet_record,
    )

    rows = ctx.pop("dwda_calc_sheet", None) or []
    if not isinstance(rows, list) or not rows or not isinstance(rows[0], dict):
        return
    ingested = flat_calc_row_from_sheet_record(rows[0])
    for nk, v in ingested.items():
        if nk not in ctx or not _s(ctx.get(nk)):
            ctx[nk] = v
    if ingested:
        ctx["_ecoventure_ingested"] = ingested
        prov = cell_contract_provenance()
        ctx["_ecoventure_contract_version"] = prov["contract_version"]
        ctx["_ecoventure_workbook_template_id"] = prov["workbook_template_id"]


def _cell_str(v: Any) -> str:
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return ""
    s = str(v).strip()
    # Mitigate formula injection when values are opened in Excel from Word tables
    if s and s[0] in "=+-@\t\r":
        s = "'" + s
    if len(s) > 32_768:
        s = s[:32_768]
    return s


def _truthy_exceedance(val: Any) -> bool:
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return False
    s = str(val).strip().lower()
    return s in ("y", "yes", "true", "1", "x", "exceedance", "exc")


def _numeric_compare_exceeds(result: Any, criteria: Any) -> bool:
    try:
        r = float(result)
        c = float(criteria)
    except (TypeError, ValueError):
        return False
    return r > c


def _rich_result_plain(result: Any, exceed: bool) -> RichText | str:
    text = _cell_str(result) if result is not None else ""
    if not exceed:
        return text
    rt = RichText()
    rt.add(text, bold=True, color="FF0000")
    return rt


_TABLE_LINK_COLUMNS = (
    "project_id",
    "project_number",
    "site_name",
    "uwi",
    "well_name",
)


def _project_row_is_blank(row: pd.Series) -> bool:
    for col in row.index:
        if _cell_str(row[col]):
            return False
    return True


def _project_row_is_duplicate_header(row: pd.Series, columns: pd.Index) -> bool:
    """Skip a data row that repeats column headers (headers pasted in row 2)."""
    match_count = 0
    filled = 0
    for col in columns:
        val = _cell_str(row[col])
        if not val:
            continue
        filled += 1
        if _norm_key(val) == _norm_key(str(col)):
            match_count += 1
    return filled > 0 and match_count >= max(2, int(filled * 0.5))


def _excel_row_number(dataframe_index: int) -> int:
    """Excel row number for a ProjectData data row (row 1 = headers)."""
    return dataframe_index + 2


def _project_row_to_dict_at(df: pd.DataFrame, index: int) -> dict[str, Any]:
    if df.empty or index < 0 or index >= len(df):
        return {}
    row = df.iloc[index]
    out: dict[str, Any] = {}
    for col in df.columns:
        key = _norm_key(col)
        if not key:
            continue
        out[key] = _cell_str(row[col])
    return out


def _project_rows_from_df(df: pd.DataFrame) -> list[dict[str, Any]]:
    """
    ProjectData data rows only.

    Excel layout: row 1 = column headers (``header=0`` for pandas);
    row 2 onward = one report per non-blank row.
    """
    rows: list[dict[str, Any]] = []
    for i in range(len(df)):
        row = df.iloc[i]
        if _project_row_is_blank(row):
            continue
        if _project_row_is_duplicate_header(row, df.columns):
            continue
        rec = _project_row_to_dict_at(df, i)
        rec["_excel_row_number"] = _excel_row_number(i)
        rows.append(rec)
    return rows


def _project_row_to_dict(df: pd.DataFrame) -> dict[str, Any]:
    rows = _project_rows_from_df(df)
    return rows[0] if rows else {}


def _filter_records_for_project(
    records: list[dict[str, Any]], project: dict[str, Any]
) -> list[dict[str, Any]]:
    """When table sheets include a link column, keep rows matching this project."""
    if not records:
        return []
    for col in _TABLE_LINK_COLUMNS:
        project_val = _s(project.get(col))
        if not project_val:
            continue
        if col not in records[0]:
            continue
        matched = [r for r in records if _s(r.get(col)) == project_val]
        if matched:
            return matched
    return records


@dataclass
class BatchReportResult:
    """One rendered report from a ProjectData row."""

    project_row_index: int
    excel_row_number: int
    docx_bytes: bytes
    warnings: list[str]
    context: dict[str, Any]
    record: Any
    filename: str
    row_label: str = ""
    appendices: list[Any] = field(default_factory=list)


def _lab_frame_to_records(df: pd.DataFrame) -> list[dict[str, Any]]:
    if df.empty:
        return []
    if len(df) > MAX_LAB_ROWS:
        df = df.head(MAX_LAB_ROWS)
    rows: list[dict[str, Any]] = []
    col_map = {_norm_key(c): i for i, c in enumerate(df.columns)}

    def get_at(row: tuple[Any, ...], *names: str) -> Any:
        for n in names:
            idx = col_map.get(n)
            if idx is not None:
                return row[idx]
        return None

    for row in df.itertuples(index=False, name=None):
        analyte = get_at(row, "analyte", "parameter", "constituent")
        result = get_at(row, "result", "value")
        unit = get_at(row, "unit", "units")
        criteria = get_at(
            row,
            "criteria",
            "standard",
            "screening_level",
            "tier1_limit",
            "background_limit",
            "guideline",
        )
        exc_col = get_at(row, "exceedance", "exceeds", "flag")

        exceed = _truthy_exceedance(exc_col) or _numeric_compare_exceeds(
            result, criteria
        )
        rec: dict[str, Any] = {}
        for i, c in enumerate(df.columns):
            k = _norm_key(c)
            if k:
                rec[k] = _cell_str(row[i])
        rec["analyte"] = _cell_str(analyte)
        rec["result"] = _cell_str(result)
        rec["unit"] = _cell_str(unit)
        rec["criteria"] = _cell_str(criteria)
        rec["exceedance_flag"] = "Yes" if exceed else "No"
        rec["result_plain"] = _cell_str(result)
        rec["result_display"] = _rich_result_plain(result, exceed)
        rows.append(rec)
    return rows


def _dataframe_to_records(df: pd.DataFrame) -> list[dict[str, Any]]:
    """Generic sheet → list of dicts (normalized keys, string cells)."""
    if df.empty:
        return []
    if len(df) > MAX_LAB_ROWS:
        df = df.head(MAX_LAB_ROWS)
    cols = [c for c in df.columns if _norm_key(c)]
    keys = [_norm_key(c) for c in cols]
    return [
        {k: _cell_str(v) for k, v in zip(keys, row)}
        for row in df[cols].itertuples(index=False, name=None)
    ]


def collect_template_root_vars(template_bytes: bytes) -> set[str]:
    """Root-level ``{{ var }}`` names from a Word template (single ZIP scan)."""
    from template_tools import scan_template

    return set(scan_template(template_bytes).root_vars)


def _runtime_cache_key(runtime: ReportRuntimeConfig) -> tuple[Any, ...]:
    return (
        runtime.report_type,
        runtime.primary_sheet,
        runtime.require_lab_sheet,
        runtime.lab_loop_variable,
        tuple(sorted(runtime.sheet_to_loop.items())),
        tuple(runtime.required_sheets),
        tuple(sorted(runtime.template_loops)),
    )


def _labels_from_project_rows(project_rows: list[dict[str, Any]]) -> list[str]:
    labels: list[str] = []
    for i, row in enumerate(project_rows):
        excel_row = int(row.get("_excel_row_number") or _excel_row_number(i))
        site = _s(row.get("site_name") or row.get("well_name"))
        client = _s(row.get("client_name"))
        proj = _s(row.get("project_number"))
        parts = [p for p in (site, client, proj) if p]
        labels.append(
            f"Excel row {excel_row}: {' — '.join(parts)}"
            if parts
            else f"Excel row {excel_row}"
        )
    return labels


class ReportEngine:
    """Load Excel + template bytes, build context, render docx."""

    def __init__(self, excel_bytes: bytes, template_bytes: bytes) -> None:
        if not validation_bypass_enabled():
            validate_excel_upload(excel_bytes)
            validate_template_upload(template_bytes)
        self.excel_bytes = excel_bytes
        self.template_bytes = template_bytes
        self._root_vars_cache: set[str] | None = None
        self._template_loops_cache: set[str] | None = None
        self._excel_cache_key: tuple[Any, ...] | None = None
        self._excel_cache: tuple[list[dict[str, Any]], dict[str, list]] | None = None
        self._excel_meta_cache: tuple[list[str], dict[str, str]] | None = None
        self._config_cache: dict[tuple[tuple[str, str], ...], ReportRuntimeConfig] = {}

    def _ensure_template_scan(self) -> None:
        """One ZIP pass for root vars and table loop names."""
        if self._root_vars_cache is not None and self._template_loops_cache is not None:
            return
        from report_profile import loops_from_block_tags
        from template_tools import scan_template

        scan = scan_template(self.template_bytes)
        self._root_vars_cache = scan.root_vars
        self._template_loops_cache = loops_from_block_tags(scan.block_tags)

    def template_root_vars(self) -> set[str]:
        self._ensure_template_scan()
        assert self._root_vars_cache is not None
        return self._root_vars_cache

    def _get_excel_meta(self) -> tuple[list[str], dict[str, str]]:
        if self._excel_meta_cache is None:
            from report_profile import read_excel_meta

            self._excel_meta_cache = read_excel_meta(self.excel_bytes)
        return self._excel_meta_cache

    def seed_template_scan(
        self, root_vars: set[str], template_loops: set[str] | None = None
    ) -> None:
        """Reuse a prior template ZIP scan (e.g. from preflight)."""
        self._root_vars_cache = root_vars
        if template_loops is not None:
            self._template_loops_cache = template_loops

    def resolve_config(self, meta: dict[str, str] | None) -> ReportRuntimeConfig:
        meta_key = tuple(sorted(sanitize_meta(meta).items()))
        cached = self._config_cache.get(meta_key)
        if cached is not None:
            return cached
        self._ensure_template_scan()
        runtime = resolve_report_config(
            self.excel_bytes,
            self.template_bytes,
            meta,
            template_loops=self._template_loops_cache,
            excel_meta=self._get_excel_meta(),
        )
        self._config_cache[meta_key] = runtime
        return runtime

    def coverage(
        self,
        meta: dict[str, str] | None,
        *,
        context: dict[str, Any] | None = None,
    ) -> "TemplateCoverage":
        from template_tools import TemplateCoverage

        if context is None:
            ctx = self.build_context(meta)
        else:
            ctx = context
        needed = self.template_root_vars()
        ctx_keys = list_keys_from_context(ctx)
        counts = table_row_counts(ctx)
        return TemplateCoverage(
            template_vars=needed,
            context_keys=ctx_keys,
            matched=sorted(needed & ctx_keys),
            missing_in_data=sorted(needed - ctx_keys),
            unused_in_template=sorted(ctx_keys - needed),
            lab_row_count=counts.get("lab_results", 0),
            drilling_waste_row_count=counts.get("drilling_waste", 0),
            storage_tanks_row_count=counts.get("storage_tanks", 0),
            table_row_counts=counts,
        )

    def _get_parsed_excel(
        self, runtime: ReportRuntimeConfig
    ) -> tuple[list[dict[str, Any]], dict[str, list]]:
        key = _runtime_cache_key(runtime)
        if self._excel_cache_key == key and self._excel_cache is not None:
            return self._excel_cache
        parsed = self._read_excel(runtime)
        self._excel_cache_key = key
        self._excel_cache = parsed
        return parsed

    def _read_excel(
        self, runtime: ReportRuntimeConfig
    ) -> tuple[list[dict[str, Any]], dict[str, list]]:
        bio = io.BytesIO(self.excel_bytes)
        with pd.ExcelFile(bio, engine="openpyxl") as xl:
            names = xl.sheet_names
            for req in runtime.required_sheets:
                if req not in names:
                    raise ValueError(
                        f"Report type '{runtime.report_type}' requires sheet "
                        f"'{req}'. Found: {names}"
                    )
            if runtime.require_lab_sheet and LAB_SHEET not in names:
                raise ValueError(
                    f"Report type '{runtime.report_type}' requires sheet "
                    f"'{LAB_SHEET}'. Found: {names}"
                )
            primary = runtime.primary_sheet
            if primary not in names:
                raise ValueError(
                    f"Missing primary sheet '{primary}'. Found: {names}"
                )
            project_df = xl.parse(primary, header=0)
            if project_df.empty:
                raise ValueError(
                    f"Sheet '{primary}' has no data rows "
                    "(row 1 = headers, row 2+ = values)."
                )
            if len(project_df.columns) > MAX_PROJECT_COLUMNS:
                project_df = project_df.iloc[:, :MAX_PROJECT_COLUMNS]
            project_rows = _project_rows_from_df(project_df)
            if not project_rows:
                raise ValueError(
                    f"Sheet '{primary}' has no data rows "
                    "(row 1 = headers, row 2+ = values)."
                )
            if len(project_rows) > MAX_PROJECT_ROWS:
                project_rows = project_rows[:MAX_PROJECT_ROWS]

            lists: dict[str, list] = {}
            lab_var = runtime.lab_loop_variable or "lab_results"
            for sheet_name, loop_var in runtime.sheet_to_loop.items():
                if sheet_name == primary or sheet_name not in names:
                    continue
                df = xl.parse(sheet_name, header=0)
                if (
                    loop_var
                    in (
                        lab_var,
                        "lab_results",
                        "groundwater_results",
                        "confirmatory_sampling",
                    )
                    or sheet_name
                    in (LAB_SHEET, GROUNDWATER_LAB_SHEET, "ConfirmatorySampling")
                ):
                    lists[loop_var] = _lab_frame_to_records(df)
                else:
                    lists[loop_var] = _dataframe_to_records(df)

            for loop_var in runtime.template_loops:
                lists.setdefault(loop_var, [])

            if runtime.require_lab_sheet and LAB_SHEET in names:
                lists.setdefault(
                    lab_var,
                    _lab_frame_to_records(xl.parse(LAB_SHEET, header=0)),
                )

        return project_rows, lists

    def project_row_count(self, meta: dict[str, str] | None = None) -> int:
        runtime = self.resolve_config(meta)
        project_rows, _ = self._get_parsed_excel(runtime)
        return len(project_rows)

    def project_row_labels(self, meta: dict[str, str] | None = None) -> list[str]:
        runtime = self.resolve_config(meta)
        project_rows, _ = self._get_parsed_excel(runtime)
        return _labels_from_project_rows(project_rows)

    def build_context(
        self,
        meta: dict[str, str] | None,
        *,
        project_row_index: int = 0,
        parsed_excel: tuple[list[dict[str, Any]], dict[str, list]] | None = None,
        appendix_labels_present: set[str] | None = None,
    ) -> dict[str, Any]:
        meta = sanitize_meta(meta)
        runtime = self.resolve_config(meta)
        if parsed_excel is not None:
            project_rows, list_data = parsed_excel
        else:
            project_rows, list_data = self._get_parsed_excel(runtime)
        if project_row_index < 0 or project_row_index >= len(project_rows):
            raise ValueError(
                f"ProjectData data row {project_row_index + 1} of {len(project_rows)} "
                f"is out of range (use Excel rows 2–{len(project_rows) + 1})."
            )
        project = dict(project_rows[project_row_index])
        excel_row = int(project.pop("_excel_row_number", _excel_row_number(project_row_index)))
        ctx: dict[str, Any] = {**project}
        ctx["_project_row_index"] = project_row_index
        ctx["_excel_row_number"] = excel_row
        ctx["_project_row_count"] = len(project_rows)
        for k, v in meta.items():
            nk = _norm_key(k)
            if nk:
                ctx[nk] = v if v is not None else ""
        from phrase_resolver import apply_phrase_resolution

        phrase_warnings = apply_phrase_resolution(
            ctx, project, self.excel_bytes, meta=meta
        )
        if phrase_warnings:
            ctx["_phrase_warnings"] = phrase_warnings
        for loop_var, rows in list_data.items():
            ctx[loop_var] = _filter_records_for_project(rows, project)
        for loop_var in runtime.template_loops:
            ctx.setdefault(loop_var, [])
        for legacy in (
            "lab_results",
            "drilling_waste",
            "storage_tanks",
            "apecs",
            "monitoring_wells",
            "water_levels",
            "groundwater_results",
            "field_events",
            "reclamation_tasks",
            "soil_placement",
            "vegetation",
            "remediation_objectives",
            "treatment_events",
            "confirmatory_sampling",
            "waste_manifests",
            "sample_locations",
            "dwda_checklist",
            "dwda_calc_sheet",
        ):
            ctx.setdefault(legacy, [])
        _merge_dwda_calc_sheet(ctx)
        ctx["_report_type"] = runtime.report_type
        phase = str(ctx.get("report_phase", "")).strip()
        if runtime.narrative_profile == "groundwater_monitoring":
            from groundwater_narrative import (
                build_groundwater_executive_summary,
                enrich_groundwater_context,
            )

            enrich_groundwater_context(ctx)
            if not _s(ctx.get("executive_summary")):
                ctx["executive_summary"] = build_groundwater_executive_summary(ctx)
                ctx["_executive_summary_auto_generated"] = True
        elif runtime.narrative_profile == "phase2":
            from phase2_narrative import (
                build_phase2_executive_summary,
                enrich_phase2_context,
            )

            enrich_phase2_context(ctx)
            if not _s(ctx.get("executive_summary")):
                ctx["executive_summary"] = build_phase2_executive_summary(ctx)
                ctx["_executive_summary_auto_generated"] = True
        elif runtime.narrative_profile == "remediation":
            from remediation_narrative import (
                build_remediation_executive_summary,
                enrich_remediation_context,
            )

            enrich_remediation_context(ctx)
            if not _s(ctx.get("executive_summary")):
                ctx["executive_summary"] = build_remediation_executive_summary(ctx)
                ctx["_executive_summary_auto_generated"] = True
        elif (
            runtime.narrative_profile == "phase1_alberta"
            and phase == "Phase 1"
        ):
            from phase1_decision import enrich_phase1_alberta_context

            ctx = enrich_phase1_alberta_context(
                ctx,
                meta,
                appendix_labels_present=appendix_labels_present,
                report_type=runtime.report_type,
            )
            if not _s(ctx.get("executive_summary")):
                ctx["executive_summary"] = build_phase1_executive_summary(ctx)
                ctx["_executive_summary_auto_generated"] = True
        return ctx

    def missing_template_vars(
        self, context: dict[str, Any]
    ) -> list[str]:
        needed = self.template_root_vars()
        missing: list[str] = []
        for name in sorted(needed):
            if name not in context:
                missing.append(name)
        return missing

    def dry_run(
        self,
        meta: dict[str, str] | None = None,
        *,
        excel_filename: str = "",
        template_filename: str = "",
        project_row_index: int = 0,
    ) -> tuple[dict[str, Any], list[str], "GenerationRecord"]:
        """
        Build context and manifest without rendering Word (preview / QA pattern).
        Does not fill missing template variables with empty strings.
        """
        from field_validation import contract_warnings
        from provenance import GenerationRecord, build_generation_record

        context = self.build_context(meta, project_row_index=project_row_index)
        auto_exec = context.pop("_executive_summary_auto_generated", False)
        phrase_warnings = context.pop("_phrase_warnings", [])
        row_count = int(context.pop("_project_row_count", 1))
        context.pop("_project_row_index", None)
        excel_row = int(context.pop("_excel_row_number", _excel_row_number(project_row_index)))
        context, clamp_warnings = clamp_context(context)
        warnings: list[str] = list(clamp_warnings) + list(phrase_warnings)
        if row_count > 1:
            warnings.append(
                f"ProjectData has {row_count} site row(s) (Excel rows 2+); "
                f"this preview uses Excel row {excel_row} only."
            )
        if auto_exec:
            warnings.append(
                "Executive summary auto-generated from ProjectData (Signum-style structure)."
            )
        missing = self.missing_template_vars(context)
        for m in missing:
            warnings.append(
                f"Template uses '{{{{ {m} }}}}' but no Excel/sidebar value yet."
            )
        warnings.extend(
            contract_warnings(
                context,
                report_phase=str((meta or {}).get("report_phase", "")),
                report_type=str(context.get("_report_type", "")),
            )
        )
        try:
            coverage = self.coverage(meta, context=context)
        except (ValueError, SecurityError) as e:
            coverage = None
            warnings.append(str(e))
        record = build_generation_record(
            excel_bytes=self.excel_bytes,
            template_bytes=self.template_bytes,
            meta=meta,
            coverage=coverage,
            warnings=warnings,
            missing_variables=missing,
            excel_filename=excel_filename,
            template_filename=template_filename,
            dry_run=True,
            template_source_format=str((meta or {}).get("template_source_format", "")),
            context=context,
        )
        from security import strip_internal_context_keys

        return strip_internal_context_keys(context), warnings, record

    def render(
        self, meta: dict[str, str] | None = None,
        *,
        excel_filename: str = "",
        template_filename: str = "",
        project_row_index: int = 0,
        parsed_excel: tuple[list[dict[str, Any]], dict[str, list]] | None = None,
        include_coverage: bool = True,
        appendix_labels_present: set[str] | None = None,
    ) -> tuple[bytes, list[str], dict[str, Any], "GenerationRecord"]:
        """
        Returns (docx_bytes, warnings, context).
        Warnings include template variables not supplied by Excel/meta
        (filled with empty string for render).
        """
        context = self.build_context(
            meta,
            project_row_index=project_row_index,
            parsed_excel=parsed_excel,
            appendix_labels_present=appendix_labels_present,
        )
        auto_exec = context.pop("_executive_summary_auto_generated", False)
        phrase_warnings = context.pop("_phrase_warnings", [])
        context.pop("_project_row_count", None)
        context.pop("_project_row_index", None)
        context.pop("_excel_row_number", None)
        context, clamp_warnings = clamp_context(context)
        warnings: list[str] = list(clamp_warnings) + list(phrase_warnings)
        if auto_exec:
            warnings.append(
                "Executive summary auto-generated from ProjectData (Signum-style structure). "
                "Review before client delivery."
            )
        missing_vars = self.missing_template_vars(context)
        for m in missing_vars:
            warnings.append(
                f"Template uses '{{{{ {m} }}}}' but no Excel/sidebar value; "
                "rendering with empty string."
            )
            context[m] = ""

        render_ctx = {
            k: v for k, v in context.items() if not str(k).startswith("_")
        }

        tpl_bio = io.BytesIO(self.template_bytes)
        doc = DocxTemplate(tpl_bio)
        env = SandboxedEnvironment(undefined=StrictUndefined, autoescape=False)
        try:
            doc.render(render_ctx, jinja_env=env)
        except TemplateError as e:
            raise ValueError(
                "Template rendering failed. Check Jinja2 tags and table loops "
                "in your Word template."
            ) from e
        out = io.BytesIO()
        doc.save(out)
        docx_bytes = out.getvalue()
        validate_rendered_output(docx_bytes)
        from provenance import build_generation_record

        coverage = None
        if include_coverage:
            try:
                coverage = self.coverage(meta, context=context)
            except (ValueError, SecurityError):
                coverage = None
        record = build_generation_record(
            excel_bytes=self.excel_bytes,
            template_bytes=self.template_bytes,
            meta=meta,
            coverage=coverage,
            warnings=warnings,
            missing_variables=missing_vars,
            output_bytes=docx_bytes,
            excel_filename=excel_filename,
            template_filename=template_filename,
            dry_run=False,
            template_source_format=str((meta or {}).get("template_source_format", "")),
            context=context,
        )
        return docx_bytes, warnings, context, record

    def render_batch(
        self,
        meta: dict[str, str] | None = None,
        *,
        excel_filename: str = "",
        template_filename: str = "",
        appendix_labels_present: set[str] | None = None,
    ) -> list[BatchReportResult]:
        """Render one .docx per non-blank ProjectData row."""
        runtime = self.resolve_config(meta)
        project_rows, list_data = self._get_parsed_excel(runtime)
        n = len(project_rows)
        if n > MAX_BATCH_REPORTS:
            raise ValueError(
                f"Too many ProjectData rows ({n}). Maximum batch size is "
                f"{MAX_BATCH_REPORTS} reports per run."
            )
        labels = _labels_from_project_rows(project_rows)
        filtered_per_row = [
            {
                loop_var: _filter_records_for_project(rows, project_rows[i])
                for loop_var, rows in list_data.items()
            }
            for i in range(n)
        ]
        results: list[BatchReportResult] = []
        for i in range(n):
            excel_row = int(
                project_rows[i].get("_excel_row_number", _excel_row_number(i))
            )
            docx_bytes, warnings, context, record = self.render(
                meta,
                excel_filename=excel_filename,
                template_filename=template_filename,
                project_row_index=i,
                parsed_excel=(project_rows, filtered_per_row[i]),
                include_coverage=(i == n - 1),
                appendix_labels_present=appendix_labels_present,
            )
            filename = suggested_download_name(
                context, meta or {}, project_row_index=i, batch_size=n
            )
            record.output_filename = filename
            results.append(
                BatchReportResult(
                    project_row_index=i,
                    excel_row_number=excel_row,
                    docx_bytes=docx_bytes,
                    warnings=list(warnings),
                    context=context,
                    record=record,
                    filename=filename,
                    row_label=labels[i] if i < len(labels) else f"Excel row {excel_row}",
                )
            )
        return results


def suggested_download_name(
    context: dict[str, Any],
    meta: dict[str, str],
    *,
    project_row_index: int | None = None,
    batch_size: int = 1,
) -> str:
    """Build a safe .docx filename from site/phase/date (unique per batch row)."""
    site = (
        str(context.get("site_name") or context.get("client_name") or "")
        .strip()
    )
    proj = str(context.get("project_number") or "").strip()
    phase = str(meta.get("report_phase") or "ESA").strip().replace(" ", "_")
    date = str(meta.get("date_of_issue") or context.get("report_month_year") or "")[:10]
    base = site or proj or "ESA"
    safe = re.sub(r"[^\w\-]+", "_", base).strip("_") or "ESA"
    parts = [safe]
    if proj and proj not in site:
        parts.append(re.sub(r"[^\w\-]+", "_", proj).strip("_"))
    parts.append(phase)
    if date:
        parts.append(re.sub(r"[^\w\-]+", "_", date).strip("_"))
    if batch_size > 1 and project_row_index is not None:
        parts.append(f"row{project_row_index + 1}")
    return sanitize_download_filename("_".join(parts) + ".docx")


def generate_sample_excel(path: str) -> None:
    """Write a minimal valid workbook (used by scripts/create_samples.py)."""
    project = pd.DataFrame(
        [
            {
                "site_name": "123 Example Road",
                "client_name": "Demo Client Ltd.",
                "project_number": "ESA-2026-001",
                "address": "Toronto, ON",
            }
        ]
    )
    lab = pd.DataFrame(
        [
            {
                "Analyte": "Benzene",
                "Result": 0.8,
                "Unit": "ug/L",
                "Criteria": 5.0,
                "Exceedance": "N",
            },
            {
                "Analyte": "TCE",
                "Result": 12.5,
                "Unit": "ug/L",
                "Criteria": 5.0,
                "Exceedance": "Y",
            },
            {
                "Analyte": "pH",
                "Result": 7.2,
                "Unit": "SU",
                "Criteria": "",
                "Exceedance": "",
            },
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        lab.to_excel(w, sheet_name=LAB_SHEET, index=False)


def generate_production_excel(path: str) -> None:
    """
    Workbook aligned with typical Phase 2 ESA fields and bracket placeholders
  found in the production merge template (map to {{ jinja }} in Word).
    """
    project = pd.DataFrame(
        [
            {
                "client_name": "Client Full Name",
                "client_full_name": "Client Full Name",
                "site_name": "1XX/XX-XX-XXX-XX WXM",
                "site_address": "Site address, City, Province",
                "project_number": "22xxxxR",
                "report_title": "Phase 2 Environmental Site Assessment",
                "report_year": "2025",
                "consultant_name": "Ecoventure Inc.",
                "company": "Ecoventure Inc.",
                "company_address": "Company address line",
                "keywords": "Phase 2 ESA, environmental assessment",
                "address": "Site address, City, Province",
                "lab_name": "Laboratory Name",
            }
        ]
    )
    lab = pd.DataFrame(
        [
            {
                "Analyte": "Benzene",
                "Result": 0.8,
                "Unit": "ug/L",
                "Criteria": 5.0,
                "Exceedance": "N",
            },
            {
                "Analyte": "TCE",
                "Result": 12.5,
                "Unit": "ug/L",
                "Criteria": 5.0,
                "Exceedance": "Y",
            },
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        lab.to_excel(w, sheet_name=LAB_SHEET, index=False)


def generate_sample_template_docx(path: str) -> None:
    """Minimal docxtpl template with ProjectData vars + lab_results table."""
    doc = Document()
    doc.add_heading("Phase 2 ESA — Lab Summary (Sample)", level=0)
    doc.add_paragraph("Site: {{ site_name }}")
    doc.add_paragraph("Client: {{ client_name }}")
    doc.add_paragraph("Project No.: {{ project_number }}")
    doc.add_paragraph("Address: {{ address }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }}")
    doc.add_paragraph("Date of issue: {{ date_of_issue }}")

    doc.add_paragraph("")
    doc.add_paragraph("Laboratory results:")

    # docxtpl table-row loop: static header, then for / data / endfor rows
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    hdr = ["Analyte", "Result", "Unit", "Exceedance"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].merge(table.rows[1].cells[3])
    table.rows[1].cells[0].text = "{%tr for item in lab_results %}"
    table.rows[2].cells[0].text = "{{ item.analyte }}"
    table.rows[2].cells[1].text = "{{ item.result_display }}"
    table.rows[2].cells[2].text = "{{ item.unit }}"
    table.rows[2].cells[3].text = "{{ item.exceedance_flag }}"
    table.rows[3].cells[0].merge(table.rows[3].cells[3])
    table.rows[3].cells[0].text = "{%tr endfor %}"

    doc.save(path)


def _add_lab_results_table(doc: Document) -> None:
    """docxtpl table-row loop: static header, for/data/endfor rows."""
    doc.add_paragraph("")
    doc.add_paragraph("Laboratory results:")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    for i, h in enumerate(["Analyte", "Result", "Unit", "Exceedance"]):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].merge(table.rows[1].cells[3])
    table.rows[1].cells[0].text = "{%tr for item in lab_results %}"
    table.rows[2].cells[0].text = "{{ item.analyte }}"
    table.rows[2].cells[1].text = "{{ item.result_display }}"
    table.rows[2].cells[2].text = "{{ item.unit }}"
    table.rows[2].cells[3].text = "{{ item.exceedance_flag }}"
    table.rows[3].cells[0].merge(table.rows[3].cells[3])
    table.rows[3].cells[0].text = "{%tr endfor %}"


def generate_production_starter_template_docx(path: str) -> None:
    """
    Tagged starter for production ESA reports — key fields + lab table.
    Use while converting the full merge template to Jinja2.
    """
    doc = Document()
    doc.add_heading("Phase 2 ESA — Production Starter", level=0)
    doc.add_paragraph("Client: {{ client_name }}")
    doc.add_paragraph("Site: {{ site_name }}")
    doc.add_paragraph("Address: {{ site_address }}")
    doc.add_paragraph("Project: {{ project_number }}")
    doc.add_paragraph("Company: {{ company }}")
    doc.add_paragraph("Company address: {{ company_address }}")
    doc.add_paragraph("Keywords: {{ keywords }}")
    doc.add_paragraph("Laboratory: {{ lab_name }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }}")
    doc.add_paragraph("Date of issue: {{ date_of_issue }}")
    _add_lab_results_table(doc)
    doc.save(path)


def generate_production_template_docx(path: str) -> None:
    """
    Full tagged production template aligned with PRODUCTION_TEMPLATE_GUIDE.txt
    and samples/production_data.xlsx (bracket placeholders as Jinja tags).
    """
    doc = Document()
    doc.add_heading("{{ report_title }}", level=0)
    doc.add_paragraph("Client: {{ client_full_name }}")
    doc.add_paragraph("Site / project: {{ site_name }}")
    doc.add_paragraph("Address: {{ site_address }}")
    doc.add_paragraph("Project number: {{ project_number }}")
    doc.add_paragraph("Report year: {{ report_year }}")
    doc.add_paragraph("Consultant: {{ consultant_name }}")
    doc.add_paragraph("Company: {{ company }}")
    doc.add_paragraph("Company address: {{ company_address }}")
    doc.add_paragraph("Keywords: {{ keywords }}")
    doc.add_paragraph("Laboratory: {{ lab_name }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }}")
    doc.add_paragraph("Date of issue: {{ date_of_issue }}")
    _add_lab_results_table(doc)
    doc.save(path)


# Bracket placeholders in the production merge doc → Jinja (see PRODUCTION_TEMPLATE_GUIDE.txt)
PRODUCTION_BRACKET_REPLACEMENTS: dict[str, str] = {
    "[Company]": "{{ company }}",
    "[Company Address]": "{{ company_address }}",
    "[Keywords]": "{{ keywords }}",
    "[LAB]": "{{ lab_name }}",
    "Client Full Name": "{{ client_full_name }}",
}


def _add_docx_table_loop(
    doc: Document,
    title: str,
    loop_var: str,
    headers: list[str],
    item_fields: list[str],
) -> None:
    """docxtpl table-row loop with static header row."""
    doc.add_paragraph("")
    doc.add_paragraph(title)
    ncols = len(headers)
    table = doc.add_table(rows=4, cols=ncols)
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].merge(table.rows[1].cells[ncols - 1])
    table.rows[1].cells[0].text = f"{{%tr for item in {loop_var} %}}"
    for i, field in enumerate(item_fields):
        table.rows[2].cells[i].text = f"{{{{ item.{field} }}}}"
    table.rows[3].cells[0].merge(table.rows[3].cells[ncols - 1])
    table.rows[3].cells[0].text = "{%tr endfor %}"


def generate_phase1_alberta_excel(path: str) -> None:
    """Alberta O&G Phase I workbook — Ecoventure Inc. sample (Signum-style executive summary)."""
    row = {
        "client_name": "Example Energy Ltd.",
        "consultant_name": ECOVENTURE_CONSULTANT,
        "company": ECOVENTURE_CONSULTANT,
        "well_name": "Example 4D Windy 4-4-49-4",
        "site_name": "Example 4D Windy 4-4-49-4",
        "uwi": "00/04-04-049-04W4/0",
        "project_number": "ESA-P1-2017-001",
        "report_title": "Phase I Environmental Site Assessment",
        "report_month_year": "March 2017",
        "asset_activity_type": "Oil well site — drilled and abandoned / suspended",
        "prior_reclamation_cert_number": "",
        "site_visit_date": "",
        "records_review_summary": (
            "Company well file, AER spills search, ABADATA, and historical air photos reviewed."
        ),
        "interview_operator_summary": (
            "Informational interview with former site operator regarding waste disposal and infrastructure."
        ),
        "interview_landowner_summary": "",
        "flare_pit_used": "No",
        "no_drilling_waste_on_site": "No",
        "site_visit_photo_notes": "",
        "qp_names": "Ecoventure QP (P.Eng.); Ecoventure QP (R.T.Ag.)",
        "spud_date": "15-Mar-2004",
        "cased_date": "17-Mar-2004",
        "final_drill_date": "19-Jun-2004",
        "well_depth_m": "710",
        "well_status": "suspended",
        "reentry": "Yes",
        "reentry_detail": (
            "The well was re-entered in June 2004 and drilled to a total depth of 710 metres (m), "
            "and the well is currently suspended"
        ),
        "production_fluid": "gas with water",
        "aer_waste_compliance_option": "Option 1 (AER, 2014)",
        "drilling_waste_summary": (
            "110 m3 of drilling waste (surface mud and fasthole mud) was disposed of via LWD at "
            "SW1/4 04-049-04 W4M, 35 m3 of mainhole drilling mud was landsprayed at SE-09-049-04 W4M, "
            "3 m3 of shale was landspread onsite, and 60 m3 of mainhole drilling mud was hauled to "
            "the remote site at Example 10-04-048-06 W4M"
        ),
        "phase2_drilling_waste_required": "No",
        "cuttings_volume_on_lease_m3": "55",
        "directive_050_notification_ref": "Directive 050 notification on file; tour reports reviewed",
        "dwda_salinity_pathway": "equivalent_salinity",
        "dwda_phase2_required": "No",
        "phase2_esa_required": "Yes",
        "client_phase_keyword": "Phase II",
        "site_visit_completed": "No",
        "air_photo_observations": (
            "The 2015 historical air photo shows the well centre and a possible disturbance area "
            "southeast of well centre for the containment berm and tanks"
        ),
        "investigations_recommended": (
            "well centre and the production areas be investigated"
        ),
        "infrastructure_summary": (
            "Access road, teardrop, wellhead, containment berm for production tanks"
        ),
        "spills_releases": "No",
        "conclusions_recommendations": (
            "Investigate well centre and production areas. Phase II ESA recommended."
        ),
        "drilling_waste_intro_selected": "option_1_aer",
        "site_recon_intro_selected": "not_completed",
        "phase2_recommendation_selected": "recommended",
    }
    row["executive_summary"] = build_phase1_executive_summary(row)
    from phrase_resolver import (
        PHRASE_CATALOG_SHEET,
        list_phrase_definitions,
        resolve_phrase_text,
    )

    for phrase_key, option_id in [
        ("drilling_waste_intro", row["drilling_waste_intro_selected"]),
        ("site_recon_intro", row["site_recon_intro_selected"]),
        ("phase2_recommendation", row["phase2_recommendation_selected"]),
    ]:
        text = resolve_phrase_text(phrase_key, option_id)
        if text:
            row[phrase_key] = text
    project = pd.DataFrame([row])
    waste = pd.DataFrame(
        [
            {
                "mud_type": "Gel Chem",
                "volume_m3": "110",
                "disposal_method": "LWD",
                "location": "SW1/4 04-049-04 W4M well centre",
                "disposal_type": "on-lease",
                "gps_coordinates": "51.1234,-114.5678",
                "sump_depth_m": "2.5",
                "cover_depth_m": "1.0",
                "remote_cert_number": "",
                "waste_manifest_refs": "",
                "dwda_id": "DWDA-1",
                "area_m2": "450",
                "salinity_exceedance": "pending",
            },
            {
                "mud_type": "Mainhole mud",
                "volume_m3": "35",
                "disposal_method": "landspray",
                "location": "SE-09-049-04 W4M",
                "disposal_type": "off-lease",
                "gps_coordinates": "",
                "sump_depth_m": "",
                "cover_depth_m": "",
                "remote_cert_number": "",
                "waste_manifest_refs": "Manifest 2017-042",
                "dwda_id": "",
                "area_m2": "",
                "salinity_exceedance": "",
            },
            {
                "mud_type": "Mainhole mud",
                "volume_m3": "60",
                "disposal_method": "remote haul",
                "location": "Example 10-04-048-06 W4M remote site",
                "disposal_type": "off-lease",
                "gps_coordinates": "",
                "sump_depth_m": "",
                "cover_depth_m": "",
                "remote_cert_number": "RC-2017-009",
                "waste_manifest_refs": "Manifest 2017-043",
                "dwda_id": "",
                "area_m2": "",
                "salinity_exceedance": "",
            },
        ]
    )
    dwda_checklist = pd.DataFrame(
        [
            {
                "checklist_item_id": "d050.notification",
                "response": "Yes",
                "notes": "Tour report on file",
            },
            {
                "checklist_item_id": "d050.compliance_option",
                "response": "Yes",
                "notes": "Option 1 AER 2014",
            },
            {
                "checklist_item_id": "d050.waste_summary",
                "response": "Yes",
                "notes": "",
            },
            {
                "checklist_item_id": "d050.waste_table",
                "response": "Yes",
                "notes": "Three disposal events in DrillingWaste",
            },
            {
                "checklist_item_id": "d050.cuttings_volume",
                "response": "Yes",
                "notes": "110 m3 LWD on lease",
            },
            {
                "checklist_item_id": "d050.gps_on_lease",
                "response": "Yes",
                "notes": "",
            },
            {
                "checklist_item_id": "d050.salinity_pathway",
                "response": "Yes",
                "notes": "Equivalent Salinity within DWDA-1",
            },
        ]
    )
    tanks = pd.DataFrame(
        [
            {
                "tank_type": "Above ground tank",
                "content": "Produced water",
                "location": "SE of well centre",
                "capacity_m3": "Unknown",
            },
        ]
    )
    apecs = pd.DataFrame(
        [
            {
                "apec_id": "APEC-1",
                "apec_name": "Historical flare pit",
                "location_description": "SW corner of lease",
                "concern_type": "flare_pit",
                "source_of_concern": "air_photo",
                "evidence_summary": "Air photos show former flare pit depression.",
                "source_document": "appendix_c_air_photos.pdf",
                "phase2_recommended": "Y",
                "notes": "Sample for confirmation if Phase II proceeds.",
            },
            {
                "apec_id": "APEC-2",
                "apec_name": "Produced water tank berm",
                "location_description": "SE of well centre",
                "concern_type": "storage_tank",
                "source_of_concern": "records",
                "evidence_summary": "Above-ground produced water tank noted in records.",
                "source_document": "historical_phase1.pdf",
                "phase2_recommended": "N",
                "notes": "",
            },
        ]
    )
    catalog_rows: list[dict[str, str]] = []
    for phrase_key, spec in sorted(list_phrase_definitions().items()):
        for opt in spec.get("options", []):
            catalog_rows.append(
                {
                    "phrase_key": phrase_key,
                    "option_id": str(opt.get("id", "")),
                    "text": str(opt.get("text", "")),
                }
            )
    phrase_catalog = pd.DataFrame(catalog_rows)

    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        waste.to_excel(w, sheet_name=DRILLING_WASTE_SHEET, index=False)
        dwda_checklist.to_excel(w, sheet_name=DWDA_CHECKLIST_SHEET, index=False)
        tanks.to_excel(w, sheet_name=STORAGE_TANKS_SHEET, index=False)
        apecs.to_excel(w, sheet_name=APECS_SHEET, index=False)
        if not phrase_catalog.empty:
            phrase_catalog.to_excel(w, sheet_name=PHRASE_CATALOG_SHEET, index=False)


def generate_phase1_alberta_template_docx(path: str) -> None:
    """Alberta Phase I ESA Word template — SED 002 Section 10 structure (Ecoventure)."""
    doc = Document()
    doc.add_heading("{{ report_title }}", level=0)
    doc.add_paragraph("PREPARED FOR: {{ client_name }}")
    doc.add_paragraph("{{ well_name }}")
    doc.add_paragraph("{{ uwi }}")
    doc.add_paragraph("")
    doc.add_paragraph("PREPARED BY:")
    doc.add_paragraph("{{ consultant_name }}")
    doc.add_paragraph("{{ qp_names }}")
    doc.add_paragraph("{{ report_month_year }}")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ executive_summary }}")
    doc.add_heading("SED 002 — Phase 1 ESA (Section 10)", level=1)
    doc.add_heading("10.1 Asset information", level=2)
    doc.add_paragraph("Activity type: {{ asset_activity_type }}")
    doc.add_paragraph("Prior reclamation certificate: {{ prior_reclamation_cert_number }}")
    doc.add_heading("10.2 Drilling information", level=2)
    doc.add_paragraph("Well: {{ well_name }} | UWI: {{ uwi }}")
    doc.add_paragraph("Spud: {{ spud_date }} | Cased: {{ cased_date }} | Final drill: {{ final_drill_date }}")
    doc.add_paragraph("TD: {{ well_depth_m }} m | Status: {{ well_status }} | Re-entry: {{ reentry }}")
    doc.add_paragraph("{{ reentry_detail }}")
    doc.add_paragraph("Production fluid: {{ production_fluid }}")
    doc.add_heading("10.4 Drilling waste disposal", level=2)
    doc.add_paragraph("Compliance option: {{ aer_waste_compliance_option }}")
    doc.add_paragraph("No drilling waste on site: {{ no_drilling_waste_on_site }}")
    doc.add_paragraph("{{ drilling_waste_intro }}")
    doc.add_paragraph("{{ drilling_waste_summary }}")
    _add_docx_table_loop(
        doc,
        "Drilling waste disposal events:",
        "drilling_waste",
        [
            "Mud type",
            "Volume (m3)",
            "Disposal type",
            "Method",
            "Location",
            "GPS",
            "Sump depth (m)",
            "Cover (m)",
            "Remote cert #",
            "Manifests",
        ],
        [
            "mud_type",
            "volume_m3",
            "disposal_type",
            "disposal_method",
            "location",
            "gps_coordinates",
            "sump_depth_m",
            "cover_depth_m",
            "remote_cert_number",
            "waste_manifest_refs",
        ],
    )
    doc.add_heading("10.5 Production and storage", level=2)
    doc.add_paragraph("Infrastructure: {{ infrastructure_summary }}")
    doc.add_paragraph("Flare pit used: {{ flare_pit_used }}")
    doc.add_paragraph("Spills/releases: {{ spills_releases }}")
    _add_docx_table_loop(
        doc,
        "Storage tanks:",
        "storage_tanks",
        ["Type", "Content", "Location", "Capacity (m3)"],
        ["tank_type", "content", "location", "capacity_m3"],
    )
    doc.add_heading("10.6 Site visit", level=2)
    doc.add_paragraph("{{ site_recon_intro }}")
    doc.add_paragraph("Site visit completed: {{ site_visit_completed }} | Date: {{ site_visit_date }}")
    doc.add_paragraph("{{ site_visit_photo_notes }}")
    doc.add_heading("10.7 Records review", level=2)
    doc.add_paragraph("{{ records_review_summary }}")
    doc.add_paragraph("Air photos: {{ air_photo_observations }}")
    _add_docx_table_loop(
        doc,
        "Areas of potential environmental concern (APECs):",
        "apecs",
        [
            "APEC ID",
            "Name",
            "Location",
            "Concern type",
            "Source",
            "Evidence",
            "Document",
            "Phase II?",
        ],
        [
            "apec_id",
            "apec_name",
            "location_description",
            "concern_type",
            "source_of_concern",
            "evidence_summary",
            "source_document",
            "phase2_recommended",
        ],
    )
    doc.add_heading("10.8 Interviews", level=2)
    doc.add_paragraph("Operator: {{ interview_operator_summary }}")
    doc.add_paragraph("Landowner: {{ interview_landowner_summary }}")
    doc.add_heading("Conclusions and Phase II", level=2)
    doc.add_paragraph("Phase II ESA required: {{ phase2_esa_required }}")
    doc.add_paragraph("Phase II recommended (rules): {{ phase2_recommended }}")
    doc.add_paragraph("{{ phase2_recommendation }}")
    doc.add_paragraph("{{ conclusions_recommendations }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }} | Date: {{ date_of_issue }}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Appendices A–H (AER forms, ABADATA, air photos, drilling waste checklists, "
        "survey, land title, site sketch) are attached in the deliverable package."
    )
    doc.save(path)


def generate_custom_demo_excel(path: str) -> None:
    """Minimal custom report workbook (template_driven profile demo)."""
    project = pd.DataFrame(
        [
            {
                "client_name": "Custom Client Ltd.",
                "site_name": "Demo Site 100",
                "report_title": "Custom Environmental Report",
                "summary_text": "This report uses a flexible profile and custom table sheet.",
            }
        ]
    )
    observations = pd.DataFrame(
        [
            {"observation": "Stressed vegetation near access road", "severity": "Low"},
            {"observation": "Stained soil near tank berm", "severity": "Medium"},
        ]
    )
    config = pd.DataFrame(
        [
            {"key": "report_type", "value": "template_driven"},
            {"key": "map_Observations", "value": "observations"},
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        observations.to_excel(w, sheet_name="Observations", index=False)
        config.to_excel(w, sheet_name="ReportConfig", index=False)


def generate_groundwater_monitoring_excel(path: str) -> None:
    """Ecoventure groundwater monitoring sample workbook."""
    from phrase_resolver import PHRASE_CATALOG_SHEET, list_phrase_definitions, resolve_phrase_text

    row = {
        "site_name": "Example Wellsite GW Program",
        "client_name": "Example Energy Ltd.",
        "project_number": "GW-2026-001",
        "address": "NE 1/4-1-1-1W5M, Alberta",
        "consultant_name": ECOVENTURE_CONSULTANT,
        "company": ECOVENTURE_CONSULTANT,
        "report_title": "Groundwater Monitoring Report",
        "monitoring_program": "Annual groundwater monitoring",
        "lab_name": "Accredited Environmental Laboratory",
        "hydrogeologic_setting": "",
        "executive_summary": "",
        "conclusions_recommendations": (
            "Continue annual monitoring. Investigate chloride exceedance at MW-2."
        ),
        "hydrograph_image_path": "",
        "site_map_image_path": "",
        "gw_program_intro_selected": "annual",
        "gw_sampling_methods_selected": "low_flow",
        "gw_data_usability_selected": "usable",
        "gw_recommendations_selected": "investigate",
    }
    for phrase_key, option_id in [
        ("gw_program_intro", row["gw_program_intro_selected"]),
        ("gw_sampling_methods", row["gw_sampling_methods_selected"]),
        ("gw_data_usability", row["gw_data_usability_selected"]),
        ("gw_recommendations", row["gw_recommendations_selected"]),
    ]:
        text = resolve_phrase_text(phrase_key, option_id)
        if text:
            row[phrase_key] = text
    project = pd.DataFrame([row])
    wells = pd.DataFrame(
        [
            {
                "well_id": "MW-1",
                "easting": "500123",
                "northing": "6000456",
                "ground_elevation_m": "985.2",
                "screen_top_m": "12.0",
                "screen_bottom_m": "18.0",
                "construction_notes": "2-inch PVC monitoring well",
            },
            {
                "well_id": "MW-2",
                "easting": "500145",
                "northing": "6000462",
                "ground_elevation_m": "984.8",
                "screen_top_m": "10.0",
                "screen_bottom_m": "16.0",
                "construction_notes": "2-inch PVC monitoring well",
            },
        ]
    )
    levels = pd.DataFrame(
        [
            {
                "well_id": "MW-1",
                "measurement_date": "2026-04-15",
                "depth_to_water_m": "8.45",
                "water_level_masl": "976.75",
            },
            {
                "well_id": "MW-2",
                "measurement_date": "2026-04-15",
                "depth_to_water_m": "7.90",
                "water_level_masl": "976.90",
            },
        ]
    )
    gw_lab = pd.DataFrame(
        [
            {
                "well_id": "MW-1",
                "sample_date": "2026-04-15",
                "analyte": "Chloride",
                "result": "120",
                "unit": "mg/L",
                "tier1_limit": "250",
                "Exceedance": "N",
            },
            {
                "well_id": "MW-2",
                "sample_date": "2025-04-15",
                "analyte": "Chloride",
                "result": "280",
                "unit": "mg/L",
                "tier1_limit": "250",
                "Exceedance": "Y",
            },
            {
                "well_id": "MW-2",
                "sample_date": "2026-04-15",
                "analyte": "Chloride",
                "result": "380",
                "unit": "mg/L",
                "tier1_limit": "250",
                "Exceedance": "Y",
            },
            {
                "well_id": "MW-2",
                "sample_date": "2026-04-15",
                "analyte": "Benzene",
                "result": "0.0005",
                "unit": "mg/L",
                "tier1_limit": "0.005",
                "Exceedance": "N",
            },
        ]
    )
    field_notes = pd.DataFrame(
        [
            {
                "well_id": "MW-1",
                "event_date": "2026-04-15",
                "purge_volume_l": "3",
                "field_observation": "Clear after purge",
            },
        ]
    )
    catalog_rows: list[dict[str, str]] = []
    for phrase_key, spec in sorted(list_phrase_definitions().items()):
        for opt in spec.get("options", []):
            catalog_rows.append(
                {
                    "phrase_key": phrase_key,
                    "option_id": str(opt.get("id", "")),
                    "text": str(opt.get("text", "")),
                }
            )
    phrase_catalog = pd.DataFrame(catalog_rows)
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        wells.to_excel(w, sheet_name=MONITORING_WELLS_SHEET, index=False)
        levels.to_excel(w, sheet_name=WATER_LEVELS_SHEET, index=False)
        gw_lab.to_excel(w, sheet_name=GROUNDWATER_LAB_SHEET, index=False)
        field_notes.to_excel(w, sheet_name="FieldNotes", index=False)
        if not phrase_catalog.empty:
            phrase_catalog.to_excel(w, sheet_name=PHRASE_CATALOG_SHEET, index=False)


def generate_groundwater_monitoring_template_docx(path: str) -> None:
    """Ecoventure groundwater monitoring Word template."""
    doc = Document()
    doc.add_heading("{{ report_title }}", level=0)
    doc.add_paragraph("Client: {{ client_name }}")
    doc.add_paragraph("Site: {{ site_name }}")
    doc.add_paragraph("Project: {{ project_number }}")
    doc.add_paragraph("Program: {{ monitoring_program }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }} | {{ consultant_name }}")
    doc.add_paragraph("Date of issue: {{ date_of_issue }}")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ executive_summary }}")
    doc.add_paragraph("{{ gw_program_intro }}")
    doc.add_paragraph("Wells: {{ well_count }} | Events: {{ monitoring_event_count }}")
    doc.add_paragraph("{{ exceedance_summary }}")
    doc.add_paragraph("{{ data_gap_note }}")
    doc.add_heading("Hydrogeologic Setting", level=1)
    doc.add_paragraph("{{ hydrogeologic_setting }}")
    doc.add_paragraph("{{ gw_sampling_methods }}")
    doc.add_heading("Figures", level=1)
    doc.add_paragraph("Hydrograph: {{ hydrograph_image_path }}")
    doc.add_paragraph("Site map: {{ site_map_image_path }}")
    _add_docx_table_loop(
        doc,
        "Monitoring wells:",
        "monitoring_wells",
        ["Well ID", "Easting", "Northing", "Screen top (m)", "Screen bottom (m)"],
        ["well_id", "easting", "northing", "screen_top_m", "screen_bottom_m"],
    )
    _add_docx_table_loop(
        doc,
        "Water levels:",
        "water_levels",
        ["Well ID", "Date", "Depth to water (m)", "Water level (m asl)"],
        ["well_id", "measurement_date", "depth_to_water_m", "water_level_masl"],
    )
    _add_docx_table_loop(
        doc,
        "Groundwater analytical results:",
        "groundwater_results",
        ["Well ID", "Analyte", "Result", "Unit", "Exceedance"],
        ["well_id", "analyte", "result_display", "unit", "exceedance_flag"],
    )
    doc.add_heading("Conclusions", level=1)
    doc.add_paragraph("{{ conclusions_recommendations }}")
    doc.add_paragraph("{{ gw_trend_summary }}")
    doc.add_paragraph("{{ gw_recommendations }}")
    doc.add_paragraph("{{ gw_data_usability }}")
    doc.save(path)


def generate_phase2_alberta_excel(path: str) -> None:
    """Alberta O&G Phase II workbook — Ecoventure sample with lab + sample locations."""
    row = {
        "client_name": "Example Energy Ltd.",
        "consultant_name": ECOVENTURE_CONSULTANT,
        "company": ECOVENTURE_CONSULTANT,
        "site_name": "Example 4D Windy 4-4-49-4",
        "site_address": "NE 1/4-04-049-04W4M, Alberta",
        "uwi": "00/04-04-049-04W4/0",
        "project_number": "ESA-P2-2017-001",
        "report_title": "Phase II Environmental Site Assessment",
        "investigation_scope": "Phase II ESA — well centre and production areas",
        "lab_name": "Accredited Environmental Laboratory",
        "qp_names": "Ecoventure QP (P.Eng.)",
        "conclusions_recommendations": (
            "Exceedances at well centre warrant risk assessment. "
            "No further Phase II off-lease sampling recommended at this time."
        ),
    }
    project = pd.DataFrame([row])
    lab = pd.DataFrame(
        [
            {
                "Sample ID": "BH-01-0.5",
                "Location": "Well centre",
                "Matrix": "Soil",
                "Depth_m": "0.5",
                "Analyte": "Benzene",
                "Result": 0.02,
                "Unit": "mg/kg",
                "Criteria": 0.01,
                "Exceedance": "Y",
            },
            {
                "Sample ID": "BH-01-2.0",
                "Location": "Well centre",
                "Matrix": "Soil",
                "Depth_m": "2.0",
                "Analyte": "Benzene",
                "Result": 0.005,
                "Unit": "mg/kg",
                "Criteria": 0.01,
                "Exceedance": "N",
            },
            {
                "Sample ID": "BH-02-1.0",
                "Location": "Tank berm SE",
                "Matrix": "Soil",
                "Depth_m": "1.0",
                "Analyte": "TCE",
                "Result": 0.001,
                "Unit": "mg/kg",
                "Criteria": 0.005,
                "Exceedance": "N",
            },
        ]
    )
    locations = pd.DataFrame(
        [
            {
                "sample_id": "BH-01-0.5",
                "location": "Well centre",
                "matrix": "Soil",
                "depth_m": "0.5",
                "sample_date": "2017-08-15",
            },
            {
                "sample_id": "BH-02-1.0",
                "location": "Tank berm SE",
                "matrix": "Soil",
                "depth_m": "1.0",
                "sample_date": "2017-08-15",
            },
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        lab.to_excel(w, sheet_name=LAB_SHEET, index=False)
        locations.to_excel(w, sheet_name="SampleLocations", index=False)


def generate_phase2_alberta_template_docx(path: str) -> None:
    """Phase II ESA Word template — lab results + sample locations."""
    doc = Document()
    doc.add_heading("{{ report_title }}", level=0)
    doc.add_paragraph("Client: {{ client_name }}")
    doc.add_paragraph("Site: {{ site_name }}")
    doc.add_paragraph("UWI: {{ uwi }}")
    doc.add_paragraph("Scope: {{ investigation_scope }}")
    doc.add_paragraph("Laboratory: {{ lab_name }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }} | {{ date_of_issue }}")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ executive_summary }}")
    doc.add_paragraph("{{ exceedance_summary }}")
    _add_docx_table_loop(
        doc,
        "Sample locations:",
        "sample_locations",
        ["Sample ID", "Location", "Matrix", "Depth (m)", "Date"],
        ["sample_id", "location", "matrix", "depth_m", "sample_date"],
    )
    _add_lab_results_table(doc)
    doc.add_heading("Conclusions", level=1)
    doc.add_paragraph("{{ conclusions_recommendations }}")
    doc.save(path)


def generate_phase3_remediation_excel(path: str) -> None:
    """Phase III remediation sample workbook."""
    from phrase_resolver import PHRASE_CATALOG_SHEET, list_phrase_definitions, resolve_phrase_text

    row = {
        "client_name": "Example Energy Ltd.",
        "consultant_name": ECOVENTURE_CONSULTANT,
        "company": ECOVENTURE_CONSULTANT,
        "site_name": "Example Wellsite Remediation",
        "project_number": "REM-2026-001",
        "report_title": "Remediation Progress Report",
        "rap_status": "Approved RAP — active",
        "remediation_status_selected": "ongoing",
        "confirmatory_summary_selected": "not_met",
        "closure_recommendation_selected": "defer",
        "conclusions_recommendations": (
            "Continue treatment until confirmatory objectives are met for benzene at MW-2."
        ),
    }
    for phrase_key, option_id in [
        ("remediation_status", row["remediation_status_selected"]),
        ("confirmatory_summary", row["confirmatory_summary_selected"]),
        ("closure_recommendation", row["closure_recommendation_selected"]),
    ]:
        text = resolve_phrase_text(phrase_key, option_id)
        if text:
            row[phrase_key] = text
    project = pd.DataFrame([row])
    objectives = pd.DataFrame(
        [
            {
                "media": "Groundwater",
                "parameter": "Benzene",
                "objective": "0.005 mg/L",
                "rap_reference": "RAP Table 3-1",
            },
        ]
    )
    treatments = pd.DataFrame(
        [
            {
                "event_date": "2026-03-01",
                "method": "Air sparging",
                "location": "MW-2 vicinity",
                "volume_m3": "N/A",
                "contractor": "Remediation contractor",
            },
        ]
    )
    confirmatory = pd.DataFrame(
        [
            {
                "well_id": "MW-2",
                "sample_date": "2026-04-15",
                "analyte": "Benzene",
                "result": "0.008",
                "unit": "mg/L",
                "tier1_limit": "0.005",
                "Exceedance": "Y",
            },
        ]
    )
    manifests = pd.DataFrame(
        [
            {
                "manifest_number": "WM-2026-014",
                "waste_type": "Impacted soil",
                "destination": "Approved facility",
                "volume_m3": "12",
            },
        ]
    )
    catalog_rows: list[dict[str, str]] = []
    for phrase_key, spec in sorted(list_phrase_definitions().items()):
        for opt in spec.get("options", []):
            catalog_rows.append(
                {
                    "phrase_key": phrase_key,
                    "option_id": str(opt.get("id", "")),
                    "text": str(opt.get("text", "")),
                }
            )
    phrase_catalog = pd.DataFrame(catalog_rows)
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        objectives.to_excel(w, sheet_name="RemediationObjectives", index=False)
        treatments.to_excel(w, sheet_name="TreatmentEvents", index=False)
        confirmatory.to_excel(w, sheet_name="ConfirmatorySampling", index=False)
        manifests.to_excel(w, sheet_name="WasteManifests", index=False)
        if not phrase_catalog.empty:
            phrase_catalog.to_excel(w, sheet_name=PHRASE_CATALOG_SHEET, index=False)


def generate_phase3_remediation_template_docx(path: str) -> None:
    """Phase III remediation Word template."""
    doc = Document()
    doc.add_heading("{{ report_title }}", level=0)
    doc.add_paragraph("Client: {{ client_name }}")
    doc.add_paragraph("Site: {{ site_name }}")
    doc.add_paragraph("RAP status: {{ rap_status }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }} | {{ date_of_issue }}")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ executive_summary }}")
    doc.add_paragraph("Objectives: {{ objective_count }} | Treatments: {{ treatment_event_count }}")
    doc.add_paragraph("{{ confirmatory_status }}")
    _add_docx_table_loop(
        doc,
        "Remediation objectives:",
        "remediation_objectives",
        ["Media", "Parameter", "Objective", "RAP ref"],
        ["media", "parameter", "objective", "rap_reference"],
    )
    _add_docx_table_loop(
        doc,
        "Treatment events:",
        "treatment_events",
        ["Date", "Method", "Location", "Volume", "Contractor"],
        ["event_date", "method", "location", "volume_m3", "contractor"],
    )
    _add_docx_table_loop(
        doc,
        "Confirmatory sampling:",
        "confirmatory_sampling",
        ["Well ID", "Date", "Analyte", "Result", "Unit", "Exceedance"],
        ["well_id", "sample_date", "analyte", "result_display", "unit", "exceedance_flag"],
    )
    doc.add_heading("Conclusions", level=1)
    doc.add_paragraph("{{ conclusions_recommendations }}")
    doc.add_paragraph("{{ closure_recommendation }}")
    doc.save(path)


def generate_reclamation_certificate_excel(path: str) -> None:
    """Reclamation certificate sample workbook."""
    row = {
        "client_name": "Example Energy Ltd.",
        "consultant_name": ECOVENTURE_CONSULTANT,
        "company": ECOVENTURE_CONSULTANT,
        "well_name": "Example 4D Windy 4-4-49-4",
        "site_name": "Example 4D Windy 4-4-49-4",
        "uwi": "00/04-04-049-04W4/0",
        "project_number": "RC-2026-001",
        "report_title": "Reclamation Certificate Application — Phase I ESA Summary",
        "certificate_status": "Application in preparation",
        "executive_summary": "",
        "conclusions_recommendations": (
            "Site reclamation meets equivalent land capability objectives. "
            "Recommend submission of reclamation certificate application."
        ),
        "asset_activity_type": "Oil well site — suspended",
        "aer_waste_compliance_option": "Option 1 (AER, 2014)",
        "drilling_waste_summary": "Drilling waste disposed per company records.",
        "records_review_summary": "Records reviewed.",
        "site_visit_completed": "Yes",
        "qp_names": "Ecoventure QP",
    }
    project = pd.DataFrame([row])
    tasks = pd.DataFrame(
        [
            {
                "task": "Remove infrastructure",
                "status": "Complete",
                "completion_date": "2025-10-01",
            },
            {
                "task": "Contour and revegetate",
                "status": "Complete",
                "completion_date": "2026-04-01",
            },
        ]
    )
    soil = pd.DataFrame(
        [
            {
                "source_area": "Well centre",
                "destination": "On-site placement",
                "volume_m3": "45",
                "placement_depth_m": "0.3",
            },
        ]
    )
    vegetation = pd.DataFrame(
        [
            {
                "species_mix": "Native grass seed mix",
                "application_rate": "Standard",
                "establishment_pct": "85",
            },
        ]
    )
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        project.to_excel(w, sheet_name=PROJECT_SHEET, index=False)
        tasks.to_excel(w, sheet_name="ReclamationTasks", index=False)
        soil.to_excel(w, sheet_name="SoilPlacement", index=False)
        vegetation.to_excel(w, sheet_name="Vegetation", index=False)


def generate_reclamation_certificate_template_docx(path: str) -> None:
    """Reclamation certificate Word template."""
    doc = Document()
    doc.add_heading("{{ report_title }}", level=0)
    doc.add_paragraph("Client: {{ client_name }}")
    doc.add_paragraph("{{ well_name }} | UWI: {{ uwi }}")
    doc.add_paragraph("Certificate status: {{ certificate_status }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }} | {{ date_of_issue }}")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph("{{ executive_summary }}")
    _add_docx_table_loop(
        doc,
        "Reclamation tasks:",
        "reclamation_tasks",
        ["Task", "Status", "Completion date"],
        ["task", "status", "completion_date"],
    )
    _add_docx_table_loop(
        doc,
        "Soil placement:",
        "soil_placement",
        ["Source", "Destination", "Volume (m3)", "Depth (m)"],
        ["source_area", "destination", "volume_m3", "placement_depth_m"],
    )
    _add_docx_table_loop(
        doc,
        "Vegetation:",
        "vegetation",
        ["Species mix", "Application rate", "Establishment %"],
        ["species_mix", "application_rate", "establishment_pct"],
    )
    doc.add_heading("Conclusions", level=1)
    doc.add_paragraph("{{ conclusions_recommendations }}")
    doc.save(path)


def generate_custom_demo_template_docx(path: str) -> None:
    """Word template with scalar fields + observations table loop."""
    doc = Document()
    doc.add_heading("{{ report_title }}", level=0)
    doc.add_paragraph("Client: {{ client_name }}")
    doc.add_paragraph("Site: {{ site_name }}")
    doc.add_paragraph("{{ summary_text }}")
    doc.add_paragraph("Prepared by: {{ prepared_by }} | Date: {{ date_of_issue }}")
    _add_docx_table_loop(
        doc,
        "Field observations:",
        "observations",
        ["Observation", "Severity"],
        ["observation", "severity"],
    )
    doc.save(path)
